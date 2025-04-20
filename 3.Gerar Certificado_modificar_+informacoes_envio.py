from docx import Document # Permite abrir e mexem em arquivos .docx
from docx.shared import Pt # com tamanho da fonte em "pontos", tipo 12 pt, igual no Word

from openpyxl import load_workbook #Abrir planilhas do excel .xlsx
import os # Sistema Operacional, caminho de arquivos, ainda não utilizada no código

from docx.shared import RGBColor #permite definir a cor da fonte (texto) usando o sistema de cores RGB (vermelho, verde, azul).

# pip install PyWin32
import win32com.client as win32
outlook = win32.Dispatch("outlook.application")

alunos = "Alunos.xlsx"
planilha_alunos = load_workbook(alunos) #

sheet_selecionada = planilha_alunos["Nomes"]

for linha in range(2, len(sheet_selecionada["A"])+ 1):

    word = Document("Certificado1.docx")

    estilo = word.styles["Normal"]


    nome_aluno = sheet_selecionada[f'A{linha}'].value
    dia = sheet_selecionada[f'B{linha}'].value
    mes = sheet_selecionada[f'C{linha}'].value
    ano = sheet_selecionada[f'D{linha}'].value
    curso = sheet_selecionada[f'E{linha}'].value
    instrutor = sheet_selecionada[f'F{linha}'].value
    email = sheet_selecionada[f'G{linha}'].value

    for paragrafo in word.paragraphs:
        if "@nome" in paragrafo.text:
            paragrafo.text = nome_aluno
            fonte = estilo.font
            fonte.name =  "Calibri (Corpo)"
            fonte.size = Pt(24)

        p1 = "Concluiu com sucesso o curso de"
        p2 = ", como carga horária de 20 horas, promovido pela escola de Cursos Online em"
        completo = f"{p1} {curso}{p2} {dia} de {mes} de {ano}." # é porque é no final da frase

        if "escola" in paragrafo.text:
            paragrafo.text = completo
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)
            nova_palavra = paragrafo.add_run(curso)
            nova_palavra.font.color.rgb = RGBColor(255, 0, 0)  # Muda para cor vermelho
            nova_palavra.underline = True  # Sublinhado
            nova_palavra.bold = True  # Negrito
            nova_palavra = paragrafo.add_run(completo)
            nova_palavra.font.color.rgb = RGBColor(0, 0, 0) # Muda para cor preto


        if "Instrutor" in paragrafo.text:
            paragrafo.text = instrutor + "- Instrutor"
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    certificado = "C:\\Users\\nicho\\PycharmProjects\\Enviando_Email\\Certificados\\" + nome_aluno + ".docx"

    word.save(certificado)

    primeiroNome = nome_aluno.split(None, 1)[0]
#Isso divide o texto nome_aluno em duas partes, com base no espaço em branco (None quer dizer "qualquer espaço").
# O 1 no .split(None, 1) quer dizer: "só quero dividir uma vez — no primeiro espaço que encontrar."
#[0] Isso pega a primeira parte da divisão, ou seja, só o primeiro nome
# split separa

    emailOutlook = outlook.CreateItem(0) #criar um e-mail
    emailOutlook.To = email #destinatário
    emailOutlook.Subject = "Certificado " + nome_aluno #Titulo do e-mail
    emailOutlook.HTMLBody = f"""
        <p>Boa noite {primeiroNome}.</p>
        <p>Segue seu <b>certificado.</b></p>
        <p>Atenciosamente, <p>
        <p>Nichole Furtado <p>
    """
    emailOutlook.Attachments.Add(certificado)#anexar certificado
    emailOutlook.Send()
print("Certificados gerados com sucesso!")
